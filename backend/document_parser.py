"""Document parser — extracts text from PDF, DOCX, TXT, CSV, and Markdown files."""
import os
import csv
import io
import re
import logging
import tempfile
from collections import Counter
from concurrent.futures import ThreadPoolExecutor, as_completed
from pathlib import Path

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Research-paper section type mapping
# ---------------------------------------------------------------------------

_SECTION_TYPES = {
    "abstract":     ["abstract"],
    "introduction": ["introduction", "intro"],
    "related_work": ["related work", "background", "literature", "prior work", "previous work"],
    "methodology":  ["methodology", "method", "approach", "proposed", "system",
                     "architecture", "model", "framework", "implementation"],
    "experiments":  ["experiment", "evaluation", "result", "performance",
                     "benchmark", "analysis", "discussion"],
    "conclusion":   ["conclusion", "summary", "future work", "future directions"],
    "references":   ["references", "bibliography"],
}

_STOPWORDS = frozenset([
    "the", "a", "an", "and", "or", "but", "in", "on", "at", "to", "for",
    "of", "with", "by", "from", "is", "are", "was", "were", "be", "been",
    "being", "have", "has", "had", "do", "does", "did", "will", "would",
    "could", "should", "may", "might", "can", "this", "that", "these",
    "those", "it", "its", "as", "we", "our", "their", "they", "he", "she",
    "his", "her", "not", "also", "than", "then", "if", "so", "such",
    "which", "who", "whom", "how", "when", "where", "what", "all", "each",
    "both", "more", "most", "other", "into", "through", "during", "before",
    "after", "above", "between", "out", "about", "up", "down", "no", "any",
    "i", "you", "your", "us", "them", "there", "here", "section", "page",
    "figure", "table", "et", "al", "paper", "show", "shown", "using",
])


def _detect_section_type(heading_text: str) -> str:
    """Map a heading string to a known section type label."""
    h = heading_text.lower()
    for section_type, keywords in _SECTION_TYPES.items():
        if any(kw in h for kw in keywords):
            return section_type
    return "body"


def _is_heading(line: str) -> tuple[bool, str]:
    """Return (is_heading, cleaned_heading_text) for a single line."""
    stripped = line.strip()
    if not stripped:
        return False, ""
    # Markdown headings: # / ## / ### / ####
    md = re.match(r'^#{1,4}\s+(.+)', stripped)
    if md:
        return True, md.group(1).strip()
    # Numbered sections: "1.", "2.1", "1.2.3 Title" — requires capital letter after number
    numbered = re.match(r'^(\d+\.(?:\d+\.?)*)\s{1,3}([A-Z].{2,60})$', stripped)
    if numbered:
        return True, stripped
    # ALL CAPS lines (4–80 chars, no table pipes, only alpha + spaces)
    if (stripped.isupper() and 4 <= len(stripped) <= 80
            and "|" not in stripped and stripped.replace(" ", "").isalpha()):
        return True, stripped
    return False, ""


def _extract_keywords(text: str, top_n: int = 5) -> list[str]:
    """Extract top-N keywords via term frequency, excluding stopwords."""
    words = re.findall(r'[a-z]{3,}', text.lower())
    words = [w for w in words if w not in _STOPWORDS]
    if not words:
        return []
    return [w for w, _ in Counter(words).most_common(top_n)]


def _is_table_line(line: str) -> bool:
    """Detect markdown / ASCII table rows."""
    stripped = line.strip()
    return bool(stripped) and stripped.count("|") >= 2


def _preprocess_image_for_ocr(image):
    """Enhance a PIL image to maximise Tesseract OCR accuracy.

    Steps that consistently improve Tesseract output:
    1. Grayscale  — removes colour noise that confuses the engine.
    2. Scale-up   — Tesseract is tuned for ~300 DPI; scaling narrow images
                    up to ≥2 000 px wide reliably lifts accuracy on scans.
    3. Sharpen    — recovers letter edges blurred by JPEG/downsampling artefacts.
    4. Contrast   — makes dark text crisper against light backgrounds.

    Returns a grayscale PIL Image ready for pytesseract.image_to_string().
    """
    from PIL import ImageFilter, ImageEnhance

    img = image.convert("L")   # grayscale

    w, h = img.size
    if w < 1800:
        scale = max(2.0, 1800 / w)
        img = img.resize((int(w * scale), int(h * scale)), image.LANCZOS
                         if hasattr(image, "LANCZOS") else 1)

    img = img.filter(ImageFilter.SHARPEN)

    enhancer = ImageEnhance.Contrast(img)
    img = enhancer.enhance(1.8)

    return img


def _clean_ocr_text(text: str) -> str:
    """Remove garbage lines from raw Tesseract output.

    Tesseract on low-quality scans produces single-character lines, symbol
    strings, and OCR-noise.  This function filters them out before chunking
    so that only lines with at least 30 % alphabetic characters survive.
    Additionally:
      - Collapses runs of whitespace within lines.
      - Removes exact-duplicate consecutive lines.
    """
    if not text:
        return ""

    lines = text.split("\n")
    cleaned: list[str] = []

    for raw in lines:
        line = raw.strip()
        if not line:
            continue
        if len(line) < 3:
            continue

        alpha   = sum(1 for c in line if c.isalpha())
        digit   = sum(1 for c in line if c.isdigit())
        useful  = alpha + digit
        ratio   = useful / len(line)

        # Keep lines with ≥30 % alpha+digit content
        if ratio < 0.30:
            continue

        # Collapse multiple spaces inside the line
        line = re.sub(r" {2,}", " ", line)
        cleaned.append(line)

    # Remove consecutive identical lines (OCR sometimes repeats lines)
    deduped: list[str] = []
    prev = None
    for line in cleaned:
        if line != prev:
            deduped.append(line)
        prev = line

    return "\n".join(deduped)


def _is_garbage_chunk(text: str) -> bool:
    """Return True if a chunk is too garbled to be worth indexing.

    A chunk is considered garbage when:
    - It is shorter than 20 characters, OR
    - Fewer than 25 % of its characters are alphabetic.
    """
    if not text or len(text.strip()) < 20:
        return True
    alpha = sum(1 for c in text if c.isalpha())
    return (alpha / len(text)) < 0.25


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _has_text(documents: list[dict]) -> bool:
    return any((d.get("text") or "").strip() for d in documents)


def _find_tesseract_cmd() -> str | None:
    """Find tesseract executable from PATH or common Windows install locations."""
    candidate_paths = [
        # Installed and exposed in PATH
        "tesseract",
        # Common machine-level path
        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
        r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
    ]

    for candidate in candidate_paths[1:]:
        if os.path.exists(candidate):
            return candidate

    # Keep PATH-based lookup last so explicit absolute paths win when available.
    return candidate_paths[0]


def _find_poppler_bin_dir() -> str | None:
    """Find Poppler bin dir from PATH aliases or common winget install layout."""
    # If aliases are already visible in PATH, pdf2image can resolve pdfinfo directly.
    if _command_exists("pdfinfo"):
        return None

    base = Path.home() / "AppData" / "Local" / "Microsoft" / "WinGet" / "Packages"
    if not base.exists():
        return None

    matches = sorted(base.glob("*oschwartz10612.Poppler*"), reverse=True)
    for match in matches:
        candidate = match / "poppler-25.07.0" / "Library" / "bin"
        if candidate.exists():
            return str(candidate)

        # Fallback for future versioned folders.
        for nested in match.glob("poppler-*/*/bin"):
            if nested.exists():
                return str(nested)
    return None


def _command_exists(command: str) -> bool:
    from shutil import which

    return which(command) is not None


def parse_file(file_path: str, filename: str = None) -> list[dict]:
    """Parse a file and return a list of document chunks with metadata."""
    if filename is None:
        filename = os.path.basename(file_path)

    ext = os.path.splitext(filename)[1].lower()

    try:
        if ext == ".pdf":
            return _parse_pdf(file_path, filename)
        elif ext == ".docx":
            return _parse_docx(file_path, filename)
        elif ext == ".txt":
            return _parse_txt(file_path, filename)
        elif ext == ".csv":
            return _parse_csv(file_path, filename)
        elif ext in (".md", ".markdown"):
            return _parse_markdown(file_path, filename)
        else:
            logger.warning(f"Unsupported file type: {ext}")
            return _parse_txt(file_path, filename)
    except Exception as e:
        logger.error(f"Error parsing {filename}: {e}")
        return []


def parse_bytes(content: bytes, filename: str) -> list[dict]:
    """Parse file content from bytes."""
    return parse_bytes_with_diagnostics(content, filename)["documents"]


def parse_bytes_with_diagnostics(content: bytes, filename: str) -> dict:
    """Parse file content from bytes and return extraction diagnostics."""
    ext = os.path.splitext(filename)[1].lower()

    diagnostics = {
        "success": False,
        "documents": [],
        "loader_used": None,
        "ocr_triggered": False,
        "error_code": None,
        "error_message": None,
        "status_messages": [],
    }

    try:
        if ext == ".pdf":
            return _parse_pdf_bytes_with_fallback(content, filename)
        elif ext in (".png", ".jpg", ".jpeg", ".tiff", ".tif", ".bmp", ".webp"):
            return _parse_image_bytes_with_diagnostics(content, filename)
        elif ext == ".docx":
            documents = _parse_docx_bytes(content, filename)
        elif ext == ".txt":
            documents = _parse_txt_bytes(content, filename)
        elif ext == ".csv":
            documents = _parse_csv_bytes(content, filename)
        elif ext in (".md", ".markdown"):
            documents = _parse_txt_bytes(content, filename)
        else:
            documents = _parse_txt_bytes(content, filename)

        diagnostics["documents"] = documents
        diagnostics["success"] = _has_text(documents)
        diagnostics["loader_used"] = ext.lstrip(".") or "txt"
        if diagnostics["success"]:
            diagnostics["status_messages"] = ["Extraction successful"]
        else:
            diagnostics["error_code"] = "NO_READABLE_TEXT"
            diagnostics["error_message"] = f"No readable text found in {filename}"
            diagnostics["status_messages"] = ["No readable text found"]
        return diagnostics
    except Exception as e:
        logger.error(f"Error parsing {filename}: {e}")
        diagnostics["error_code"] = "PARSER_FAILURE"
        diagnostics["error_message"] = f"Failed to parse {filename}: {e}"
        return diagnostics


# ---------------------------------------------------------------------------
# Chunking
# ---------------------------------------------------------------------------

def chunk_text(text: str, chunk_size: int = 500, overlap: int = 50) -> list[str]:
    """Split text into overlapping character-based chunks.

    Used for unstructured text (OCR output, CSV rows).
    For structured documents (PDF, DOCX, TXT), use chunk_structured() instead.
    """
    text = text.strip()
    if not text:
        return []
    if len(text) <= chunk_size:
        return [text]

    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        # Try to break at sentence boundary
        if end < len(text):
            for sep in [". ", "\n\n", "\n", " "]:
                last_sep = text[start:end].rfind(sep)
                if last_sep > chunk_size // 2:
                    end = start + last_sep + len(sep)
                    break
        chunks.append(text[start:end].strip())
        start = end - overlap
    return [c for c in chunks if c]


def chunk_structured(text: str, source: str, page: int = 1,
                     doc_type: str = "research_paper") -> list[dict]:
    """Structure-aware chunker for research papers and structured documents.

    Splits text at heading boundaries (section breaks) rather than fixed character
    windows, preserving logical document units. Each chunk includes:
      heading, section, doc_type, keywords, page, chunk_id, source, text.

    Falls back to character-based chunking for large sub-sections.
    Keeps table rows together (not split mid-table).

    Args:
        text:     Extracted text from a single page or document.
        source:   Filename / document name for the chunk_id.
        page:     Page number (for PDFs).
        doc_type: Document domain tag ("research_paper", "general", etc.)
    """
    MAX_CHUNK = 800
    OVERLAP   = 100

    lines = text.split("\n")

    # ── Pass 1: segment text into (heading, section_type, lines) triples ──
    sections: list[tuple[str, str, list[str]]] = []
    current_heading      = ""
    current_section_type = "body"
    current_lines: list[str] = []

    for line in lines:
        is_h, heading_text = _is_heading(line)
        if is_h:
            if current_lines:
                sections.append((current_heading, current_section_type, current_lines))
            current_heading      = heading_text
            current_section_type = _detect_section_type(heading_text)
            current_lines        = []
        else:
            current_lines.append(line)

    if current_lines:
        sections.append((current_heading, current_section_type, current_lines))

    # No headings detected → treat full text as one body section
    if not sections:
        sections = [("", "body", lines)]

    # ── Pass 2: build chunk dicts from each section ────────────────────────
    chunks: list[dict] = []
    chunk_counter = 0

    for heading, section_type, sec_lines in sections:
        section_text = "\n".join(sec_lines).strip()
        if not section_text:
            continue

        # Preserve tables as single chunks (avoid splitting mid-row)
        table_lines = [l for l in sec_lines if _is_table_line(l)]
        is_table = len(sec_lines) > 0 and len(table_lines) / len(sec_lines) > 0.5

        if is_table:
            keywords = _extract_keywords(section_text)
            chunks.append({
                "text":     section_text[:MAX_CHUNK * 2],
                "source":   source,
                "page":     page,
                "chunk_id": f"{source}_p{page}_c{chunk_counter}",
                "heading":  heading,
                "section":  section_type,
                "doc_type": doc_type,
                "keywords": keywords,
            })
            chunk_counter += 1
            continue

        # Small section → single chunk
        if len(section_text) <= MAX_CHUNK:
            keywords = _extract_keywords(section_text)
            chunks.append({
                "text":     section_text,
                "source":   source,
                "page":     page,
                "chunk_id": f"{source}_p{page}_c{chunk_counter}",
                "heading":  heading,
                "section":  section_type,
                "doc_type": doc_type,
                "keywords": keywords,
            })
            chunk_counter += 1
        else:
            # Large section → sub-chunk at sentence boundaries, inherit heading/section
            for sub in chunk_text(section_text, chunk_size=MAX_CHUNK, overlap=OVERLAP):
                keywords = _extract_keywords(sub)
                chunks.append({
                    "text":     sub,
                    "source":   source,
                    "page":     page,
                    "chunk_id": f"{source}_p{page}_c{chunk_counter}",
                    "heading":  heading,
                    "section":  section_type,
                    "doc_type": doc_type,
                    "keywords": keywords,
                })
                chunk_counter += 1

    return chunks


# --- Image OCR (PNG / JPG / TIFF / BMP) ---

def _parse_image_bytes_with_diagnostics(content: bytes, filename: str) -> dict:
    """OCR a raw image file and return diagnostics.
    Owner: Anirudh (Data Engineer) — direct image ingestion pipeline.
    Supports PNG, JPG, JPEG, TIFF, BMP, WEBP via pytesseract + Pillow.
    """
    diagnostics = {
        "success": False,
        "documents": [],
        "loader_used": "ImageOCR",
        "ocr_triggered": True,
        "error_code": None,
        "error_message": None,
        "status_messages": ["Image detected", "Applying OCR..."],
    }

    try:
        import pytesseract
        from PIL import Image

        tesseract_cmd = _find_tesseract_cmd()
        if tesseract_cmd and tesseract_cmd != "tesseract":
            pytesseract.pytesseract.tesseract_cmd = tesseract_cmd

        # Force a clear error early if Tesseract is missing
        try:
            pytesseract.get_tesseract_version()
        except Exception as exc:
            raise RuntimeError(
                "tesseract not found. Install Tesseract OCR and ensure tesseract.exe is in PATH."
            ) from exc

        image = Image.open(io.BytesIO(content))
        # Preprocess: grayscale + scale-up + sharpen + contrast boost
        image = _preprocess_image_for_ocr(image)

        raw_text = pytesseract.image_to_string(image, config=_TESSERACT_CONFIG) or ""
        text     = _clean_ocr_text(raw_text)
        chunks   = [c for c in chunk_text(text) if not _is_garbage_chunk(c)]

        if not chunks:
            diagnostics.update({
                "error_code": "NO_READABLE_TEXT",
                "error_message": f"No readable text found in image {filename}",
                "status_messages": diagnostics["status_messages"] + ["No readable text found"],
            })
            logger.error("Image OCR found no text | file=%s", filename)
            return diagnostics

        documents = [
            {
                "text": chunk,
                "source": filename,
                "page": 1,
                "chunk_id": f"{filename}_ocr_c{i}",
            }
            for i, chunk in enumerate(chunks)
        ]

        diagnostics.update({
            "success": True,
            "documents": documents,
            "status_messages": diagnostics["status_messages"] + ["Extraction successful"],
        })
        logger.info("Image OCR success | file=%s chunks=%d", filename, len(documents))
        return diagnostics

    except Exception as exc:
        error_text = str(exc).lower()
        is_dep_error = "tesseract" in error_text or "not found" in error_text or "pillow" in error_text

        if is_dep_error:
            diagnostics["error_code"] = "OCR_DEPENDENCY_MISSING"
            diagnostics["error_message"] = (
                "OCR engine unavailable. Install Tesseract OCR + Pillow and ensure tesseract.exe is in PATH."
            )
        else:
            diagnostics["error_code"] = "IMAGE_PARSE_FAILURE"
            diagnostics["error_message"] = f"Failed to process image {filename}: {exc}"

        diagnostics["status_messages"].append("OCR failed")
        logger.error("Image OCR failure | file=%s error=%s", filename, exc)
        return diagnostics


# --- PDF ---

def _parse_pdf(file_path: str, filename: str) -> list[dict]:
    try:
        from pypdf import PdfReader
    except ImportError:
        logger.error("pypdf not installed. Install with: pip install pypdf")
        return []

    reader = PdfReader(file_path)
    documents = []
    for page_num, page in enumerate(reader.pages, 1):
        text = page.extract_text() or ""
        if not text.strip():
            continue
        documents.extend(chunk_structured(text, filename, page=page_num))
    logger.info(f"Parsed PDF {filename}: {len(documents)} chunks from {len(reader.pages)} pages")
    return documents


def _parse_pdf_bytes(content: bytes, filename: str) -> list[dict]:
    try:
        from pypdf import PdfReader
    except ImportError:
        logger.error("pypdf not installed")
        return []

    reader = PdfReader(io.BytesIO(content))
    documents = []
    for page_num, page in enumerate(reader.pages, 1):
        text = page.extract_text() or ""
        if not text.strip():
            continue
        documents.extend(chunk_structured(text, filename, page=page_num))
    return documents


def _parse_pdf_bytes_with_fallback(content: bytes, filename: str) -> dict:
    """Parse PDF via layered fallback: PyPDFLoader -> UnstructuredPDFLoader -> OCR."""
    diagnostics = {
        "success": False,
        "documents": [],
        "loader_used": None,
        "ocr_triggered": False,
        "error_code": None,
        "error_message": None,
        "status_messages": [],
    }

    loader_errors = []

    # 1) pypdf byte parser (equivalent to PyPDFLoader text extraction stage)
    try:
        docs = _parse_pdf_bytes(content, filename)
        if _has_text(docs):
            diagnostics.update(
                {
                    "success": True,
                    "documents": docs,
                    "loader_used": "PyPDFLoader",
                    "status_messages": ["Extraction successful"],
                }
            )
            logger.info("PDF extraction success | file=%s loader=PyPDFLoader ocr_triggered=False chunks=%s", filename, len(docs))
            return diagnostics
        logger.warning("PDF extraction empty | file=%s loader=PyPDFLoader", filename)
        diagnostics["status_messages"].append("Scanned document detected")
    except Exception as exc:
        loader_errors.append(f"PyPDFLoader: {exc}")
        logger.warning("PDF extraction failure | file=%s loader=PyPDFLoader error=%s", filename, exc)

    # 2) Unstructured loader fallback (best-effort; optional dependency)
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
            tmp.write(content)
            tmp_pdf_path = tmp.name

        try:
            docs = _parse_pdf_with_unstructured_loader(tmp_pdf_path, filename)
            if _has_text(docs):
                diagnostics.update(
                    {
                        "success": True,
                        "documents": docs,
                        "loader_used": "UnstructuredPDFLoader",
                        "status_messages": ["Extraction successful"],
                    }
                )
                logger.info("PDF extraction success | file=%s loader=UnstructuredPDFLoader ocr_triggered=False chunks=%s", filename, len(docs))
                return diagnostics
            logger.warning("PDF extraction empty | file=%s loader=UnstructuredPDFLoader", filename)
        finally:
            try:
                os.unlink(tmp_pdf_path)
            except OSError:
                pass
    except Exception as exc:
        loader_errors.append(f"UnstructuredPDFLoader: {exc}")
        logger.warning("PDF extraction failure | file=%s loader=UnstructuredPDFLoader error=%s", filename, exc)

    # 3) OCR fallback
    diagnostics["ocr_triggered"] = True
    diagnostics["status_messages"] = ["Scanned document detected", "Applying OCR..."]
    try:
        docs = _parse_pdf_with_ocr(content, filename)
        if _has_text(docs):
            diagnostics.update(
                {
                    "success": True,
                    "documents": docs,
                    "loader_used": "OCR",
                    "status_messages": diagnostics["status_messages"] + ["Extraction successful"],
                }
            )
            logger.info("PDF extraction success | file=%s loader=OCR ocr_triggered=True chunks=%s", filename, len(docs))
            return diagnostics

        diagnostics.update(
            {
                "loader_used": "OCR",
                "error_code": "NO_READABLE_TEXT",
                "error_message": f"No readable text found in {filename}",
                "status_messages": diagnostics["status_messages"] + ["No readable text found"],
            }
        )
        logger.error("PDF extraction failure | file=%s loader=OCR ocr_triggered=True reason=no_readable_text", filename)
        return diagnostics
    except Exception as exc:
        loader_errors.append(f"OCR: {exc}")
        error_text = str(exc).lower()
        is_dependency_error = (
            "tesseract" in error_text
            or "poppler" in error_text
            or "pdfinfo" in error_text
            or "is not installed" in error_text
            or "not found" in error_text
        )

        if is_dependency_error:
            error_code = "OCR_DEPENDENCY_MISSING"
            error_message = (
                "OCR engine unavailable. Install Tesseract OCR and Poppler, then ensure both are in PATH."
            )
            status_messages = diagnostics["status_messages"] + ["OCR engine unavailable"]
        else:
            error_code = "CORRUPTED_OR_UNREADABLE_PDF"
            error_message = f"Failed to parse PDF {filename}. The file may be corrupted or unreadable."
            status_messages = diagnostics["status_messages"] + ["No readable text found"]

        diagnostics.update(
            {
                "loader_used": "OCR",
                "error_code": error_code,
                "error_message": error_message,
                "status_messages": status_messages,
            }
        )
        logger.error("PDF extraction failure | file=%s loader=OCR ocr_triggered=True error=%s", filename, exc)
        return diagnostics
    finally:
        if loader_errors:
            logger.info("PDF loader attempts | file=%s details=%s", filename, " | ".join(loader_errors))


def _parse_pdf_with_unstructured_loader(file_path: str, filename: str) -> list[dict]:
    from langchain_community.document_loaders import UnstructuredPDFLoader

    loader = UnstructuredPDFLoader(file_path)
    raw_docs = loader.load()

    documents = []
    for idx, doc in enumerate(raw_docs, 1):
        text = (doc.page_content or "").strip()
        if not text:
            continue
        documents.extend(chunk_structured(text, filename, page=idx))
    return documents


_TESSERACT_CONFIG = "--psm 3 --oem 3"


def _ocr_single_page(args: tuple) -> tuple[int, str]:
    """OCR one page image.  Designed to run in a ThreadPoolExecutor worker.

    Tesseract releases the GIL during its C-level processing, so true thread
    parallelism is possible here — no need for multiprocessing.

    Steps:
      1. Preprocess (grayscale + scale-up + sharpen + contrast).
      2. Run Tesseract with explicit PSM/OEM flags.
      3. Clean output: remove garbage lines, collapse spaces, deduplicate.

    Returns (page_num, cleaned_text).
    """
    import pytesseract

    page_num, image = args
    try:
        processed = _preprocess_image_for_ocr(image)
        raw  = pytesseract.image_to_string(processed, config=_TESSERACT_CONFIG) or ""
        text = _clean_ocr_text(raw)
        return page_num, text
    except Exception as exc:
        logger.warning("OCR failed for page %d: %s", page_num, exc)
        return page_num, ""


def _parse_pdf_with_ocr(content: bytes, filename: str) -> list[dict]:
    """OCR-based PDF parsing with parallel page processing and image preprocessing.

    Performance design:
    - All pages are converted to PIL images first (one call to convert_from_bytes).
    - Pages are OCR-ed in a ThreadPoolExecutor (up to 4 workers).
      Tesseract releases the GIL so threads achieve true parallelism.
    - Results are collected and sorted back into page order.

    Quality design:
    - Each image is preprocessed (grayscale, scale-up, sharpen, contrast boost)
      before being passed to Tesseract.
    - Raw OCR output is cleaned (_clean_ocr_text) to remove garbage lines.
    - Chunks with < 25 % alpha content are discarded (_is_garbage_chunk).
    """
    import pytesseract
    from pdf2image import convert_from_bytes

    tesseract_cmd = _find_tesseract_cmd()
    if tesseract_cmd and tesseract_cmd != "tesseract":
        pytesseract.pytesseract.tesseract_cmd = tesseract_cmd

    poppler_path = _find_poppler_bin_dir()

    try:
        pytesseract.get_tesseract_version()
    except Exception as exc:
        raise RuntimeError(
            "tesseract not found. Install Tesseract OCR and ensure tesseract.exe is available in PATH."
        ) from exc

    if poppler_path:
        images = convert_from_bytes(content, poppler_path=poppler_path)
    else:
        images = convert_from_bytes(content)

    if not images:
        return []

    # Parallel OCR — up to 4 threads (avoid over-subscribing on single-core machines)
    max_workers = min(4, len(images))
    page_results: dict[int, str] = {}

    with ThreadPoolExecutor(max_workers=max_workers) as pool:
        futures = {
            pool.submit(_ocr_single_page, (page_num, img)): page_num
            for page_num, img in enumerate(images, 1)
        }
        for future in as_completed(futures):
            page_num, text = future.result()
            page_results[page_num] = text

    documents = []
    for page_num in sorted(page_results):
        text = page_results[page_num]
        if not text.strip():
            continue
        for i, chunk in enumerate(chunk_text(text)):
            if _is_garbage_chunk(chunk):
                continue
            documents.append({
                "text":     chunk,
                "source":   filename,
                "page":     page_num,
                "chunk_id": f"{filename}_ocr_p{page_num}_c{i}",
            })

    return documents


# --- DOCX ---

def _parse_docx(file_path: str, filename: str) -> list[dict]:
    try:
        from docx import Document
    except ImportError:
        logger.error("python-docx not installed. Install with: pip install python-docx")
        return []

    doc = Document(file_path)
    full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    documents = chunk_structured(full_text, filename, page=1, doc_type="general")
    logger.info(f"Parsed DOCX {filename}: {len(documents)} chunks")
    return documents


def _parse_docx_bytes(content: bytes, filename: str) -> list[dict]:
    try:
        from docx import Document
    except ImportError:
        return []

    doc = Document(io.BytesIO(content))
    full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    return chunk_structured(full_text, filename, page=1, doc_type="general")


# --- TXT / Markdown ---

def _parse_txt(file_path: str, filename: str) -> list[dict]:
    with open(file_path, "r", encoding="utf-8") as f:
        text = f.read()
    return _text_to_chunks(text, filename)


def _parse_txt_bytes(content: bytes, filename: str) -> list[dict]:
    text = content.decode("utf-8", errors="ignore")
    return _text_to_chunks(text, filename)


def _parse_markdown(file_path: str, filename: str) -> list[dict]:
    return _parse_txt(file_path, filename)


def _text_to_chunks(text: str, filename: str) -> list[dict]:
    return chunk_structured(text, filename, page=1, doc_type="general")


# --- CSV ---

def _parse_csv(file_path: str, filename: str) -> list[dict]:
    with open(file_path, "r", encoding="utf-8") as f:
        return _csv_to_chunks(f, filename)


def _parse_csv_bytes(content: bytes, filename: str) -> list[dict]:
    text = content.decode("utf-8", errors="ignore")
    return _csv_to_chunks(io.StringIO(text), filename)


def _csv_to_chunks(file_obj, filename: str) -> list[dict]:
    reader = csv.reader(file_obj)
    rows = list(reader)
    if not rows:
        return []

    headers = rows[0] if rows else []
    documents = []
    for row_num, row in enumerate(rows[1:], 2):
        if headers:
            text = " | ".join(f"{h}: {v}" for h, v in zip(headers, row) if v.strip())
        else:
            text = " | ".join(row)
        if text.strip():
            documents.append({
                "text": text,
                "source": filename,
                "page": row_num,
                "chunk_id": f"{filename}_r{row_num}"
            })
    logger.info(f"Parsed CSV {filename}: {len(documents)} rows")
    return documents


# --- Supported formats ---

SUPPORTED_EXTENSIONS = {
    ".pdf", ".docx", ".txt", ".csv", ".md", ".markdown",
    # Image formats — OCR via pytesseract + Pillow
    ".png", ".jpg", ".jpeg", ".tiff", ".tif", ".bmp", ".webp",
}

def is_supported(filename: str) -> bool:
    ext = os.path.splitext(filename)[1].lower()
    return ext in SUPPORTED_EXTENSIONS
